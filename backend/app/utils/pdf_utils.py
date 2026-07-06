"""PDF and document parsing utilities."""
import base64
import io
from pathlib import Path
from typing import Optional


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document
    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text_from_image(file_path: str) -> str:
    """Extract text from an image using Tesseract OCR."""
    from PIL import Image
    import pytesseract
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    return Path(file_path).read_text(encoding="utf-8", errors="replace")


def extract_text(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        return extract_text_from_image(file_path)
    elif ext in (".txt", ".md", ".csv"):
        return extract_text_from_txt(file_path)
    else:
        # Try reading as text, fallback to empty
        try:
            return extract_text_from_txt(file_path)
        except Exception:
            return ""


def pdf_to_images_base64(file_path: str, max_pages: int = 5) -> list[dict]:
    """Convert PDF pages to base64-encoded images for vision LLM."""
    import fitz
    doc = fitz.open(file_path)
    images = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        b64 = base64.b64encode(img_bytes).decode("utf-8")
        images.append({"page": i + 1, "base64": b64, "mime_type": "image/png"})
    doc.close()
    return images


def image_to_base64(file_path: str) -> dict:
    """Convert an image file to base64."""
    from PIL import Image
    ext = Path(file_path).suffix.lower()
    mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif"}
    mime_type = mime_map.get(ext, "image/png")

    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return {"base64": b64, "mime_type": mime_type}
